from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

# AISCR Brand Colors
NAVY = RGBColor(31, 78, 120)  # #1F4E78
TURQUOISE = RGBColor(46, 196, 182)  # #2EC4B6

def set_cell_color(cell, color_hex):
    """Set background color of a table cell"""
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), color_hex)
    cell._element.get_or_add_tcPr().append(shading_elm)

def add_header_row(table, headers, color_hex):
    """Add a formatted header row to a table"""
    header_row = table.rows[0]
    for i, header in enumerate(headers):
        cell = header_row.cells[i]
        cell.text = header
        run = cell.paragraphs[0].runs[0]
        run.font.bold = True
        run.font.color.rgb = RGBColor(255, 255, 255)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_color(cell, color_hex)

# ============================================
# 1. Charter Template
# ============================================
doc1 = Document()
doc1.add_heading('AISCR GLOBAL PMO - PROJECT CHARTER', 0)
doc1.add_paragraph('Project Charter Template')

table1 = doc1.add_table(rows=7, cols=2)
table1.style = 'Light Grid Accent 1'
headers1 = ['Field', 'Description']
add_header_row(table1, headers1, '1F4E78')

data1 = [
    ['Project Name', ''],
    ['Project Manager', ''],
    ['Start Date', ''],
    ['End Date', ''],
    ['Objectives', ''],
    ['Success Criteria', '']
]

for i, row_data in enumerate(data1, 1):
    row = table1.rows[i]
    row.cells[0].text = row_data[0]
    row.cells[1].text = row_data[1]

doc1.save('Templates/01_Charter_Template.docx')

# ============================================
# 2. Scope Template
# ============================================
doc2 = Document()
doc2.add_heading('AISCR GLOBAL PMO - PROJECT SCOPE', 0)
doc2.add_paragraph('Project Scope Template')

table2 = doc2.add_table(rows=4, cols=4)
table2.style = 'Light Grid Accent 1'
headers2 = ['In Scope', 'Out of Scope', 'Constraints', 'Assumptions']
add_header_row(table2, headers2, '1F4E78')

for i in range(1, 4):
    row = table2.rows[i]
    for j in range(4):
        row.cells[j].text = ''

doc2.save('Templates/02_Scope_Template.docx')

# ============================================
# 3. WBS Template
# ============================================
doc3 = Document()
doc3.add_heading('AISCR GLOBAL PMO - WORK BREAKDOWN STRUCTURE', 0)
doc3.add_paragraph('WBS Template')

table3 = doc3.add_table(rows=10, cols=2)
table3.style = 'Light Grid Accent 1'
headers3 = ['WBS Level', 'Task']
add_header_row(table3, headers3, '1F4E78')

for i in range(1, 10):
    row = table3.rows[i]
    row.cells[0].text = ''
    row.cells[1].text = ''

doc3.save('Templates/03_WBS_Template.docx')

# ============================================
# 4. Risk Register Template
# ============================================
doc4 = Document()
doc4.add_heading('AISCR GLOBAL PMO - RISK REGISTER', 0)
doc4.add_paragraph('Risk Register Template - Score = Probability × Impact')

table4 = doc4.add_table(rows=10, cols=8)
table4.style = 'Light Grid Accent 1'
headers4 = ['Risk ID', 'Description', 'Probability', 'Impact', 'Score', 'Owner', 'Mitigation', 'Status']
add_header_row(table4, headers4, '1F4E78')

for i in range(1, 10):
    row = table4.rows[i]
    for j in range(8):
        row.cells[j].text = ''

doc4.add_paragraph('Status Options: Green, Yellow, Red', style='Intense Quote')

doc4.save('Templates/04_Risk_Register_Template.docx')

# ============================================
# 5. Stakeholders Template
# ============================================
doc5 = Document()
doc5.add_heading('AISCR GLOBAL PMO - STAKEHOLDER REGISTER', 0)
doc5.add_paragraph('Stakeholder Register Template')

table5 = doc5.add_table(rows=10, cols=5)
table5.style = 'Light Grid Accent 1'
headers5 = ['Stakeholder', 'Role', 'Influence', 'Interest', 'Engagement Strategy']
add_header_row(table5, headers5, '1F4E78')

for i in range(1, 10):
    row = table5.rows[i]
    for j in range(5):
        row.cells[j].text = ''

doc5.save('Templates/05_Stakeholders_Template.docx')

# ============================================
# 6. Timeline Template
# ============================================
doc6 = Document()
doc6.add_heading('AISCR GLOBAL PMO - PROJECT TIMELINE', 0)
doc6.add_paragraph('Timeline Template - Duration = End Date - Start Date')

table6 = doc6.add_table(rows=10, cols=5)
table6.style = 'Light Grid Accent 1'
headers6 = ['Task', 'Start Date', 'End Date', 'Duration', 'Status']
add_header_row(table6, headers6, '1F4E78')

for i in range(1, 10):
    row = table6.rows[i]
    for j in range(5):
        row.cells[j].text = ''

doc6.add_paragraph('Status Options: Not Started, In Progress, Complete', style='Intense Quote')

doc6.save('Templates/06_Timeline_Template.docx')

# ============================================
# 7. Budget Template
# ============================================
doc7 = Document()
doc7.add_heading('AISCR GLOBAL PMO - PROJECT BUDGET', 0)
doc7.add_paragraph('Budget Template - Variance = Actual Cost - Estimated Cost')

table7 = doc7.add_table(rows=10, cols=6)
table7.style = 'Light Grid Accent 1'
headers7 = ['Category', 'Description', 'Estimated Cost', 'Actual Cost', 'Variance', 'Notes']
add_header_row(table7, headers7, '1F4E78')

for i in range(1, 10):
    row = table7.rows[i]
    for j in range(6):
        row.cells[j].text = ''

doc7.save('Templates/07_Budget_Template.docx')

# ============================================
# 8. Issues Template
# ============================================
doc8 = Document()
doc8.add_heading('AISCR GLOBAL PMO - ISSUES REGISTER', 0)
doc8.add_paragraph('Issues Register Template')

table8 = doc8.add_table(rows=10, cols=7)
table8.style = 'Light Grid Accent 1'
headers8 = ['Issue ID', 'Description', 'Impact', 'Owner', 'Due Date', 'Status', 'Notes']
add_header_row(table8, headers8, '1F4E78')

for i in range(1, 10):
    row = table8.rows[i]
    for j in range(7):
        row.cells[j].text = ''

doc8.add_paragraph('Status Options: Open, In Progress, Closed', style='Intense Quote')

doc8.save('Templates/08_Issues_Template.docx')

# ============================================
# 9. Change Requests Template
# ============================================
doc9 = Document()
doc9.add_heading('AISCR GLOBAL PMO - CHANGE REQUEST FORM', 0)
doc9.add_paragraph('Change Request Template')

table9 = doc9.add_table(rows=7, cols=2)
table9.style = 'Light Grid Accent 1'
headers9 = ['Field', 'Content']
add_header_row(table9, headers9, '1F4E78')

data9 = [
    ['Request ID', ''],
    ['Requestor', ''],
    ['Date', ''],
    ['Description', ''],
    ['Impact', ''],
    ['Status', '']
]

for i, row_data in enumerate(data9, 1):
    row = table9.rows[i]
    row.cells[0].text = row_data[0]
    row.cells[1].text = row_data[1]

doc9.save('Templates/09_Change_Requests_Template.docx')

# ============================================
# 10. Lessons Learned Template
# ============================================
doc10 = Document()
doc10.add_heading('AISCR GLOBAL PMO - LESSONS LEARNED', 0)
doc10.add_paragraph('Lessons Learned Template')

table10 = doc10.add_table(rows=5, cols=2)
table10.style = 'Light Grid Accent 1'
headers10 = ['Aspect', 'Details']
add_header_row(table10, headers10, '1F4E78')

data10 = [
    ['What Went Well', ''],
    ['What Could Improve', ''],
    ['Recommendations', ''],
    ['Best Practices', '']
]

for i, row_data in enumerate(data10, 1):
    row = table10.rows[i]
    row.cells[0].text = row_data[0]
    row.cells[1].text = row_data[1]

doc10.save('Templates/10_Lessons_Learned_Template.docx')

# ============================================
# 11. Volunteers Template
# ============================================
doc11 = Document()
doc11.add_heading('AISCR GLOBAL PMO - VOLUNTEER REGISTER', 0)
doc11.add_paragraph('Volunteer Register Template')

table11 = doc11.add_table(rows=10, cols=5)
table11.style = 'Light Grid Accent 1'
headers11 = ['Volunteer', 'Skill', 'Level', 'Availability', 'Notes']
add_header_row(table11, headers11, '1F4E78')

for i in range(1, 10):
    row = table11.rows[i]
    for j in range(5):
        row.cells[j].text = ''

doc11.save('Templates/11_Volunteers_Template.docx')

# ============================================
# 12. Assignments Template
# ============================================
doc12 = Document()
doc12.add_heading('AISCR GLOBAL PMO - VOLUNTEER ASSIGNMENTS', 0)
doc12.add_paragraph('Volunteer Assignments Template')

table12 = doc12.add_table(rows=10, cols=6)
table12.style = 'Light Grid Accent 1'
headers12 = ['Volunteer', 'Project', 'Role', 'Start Date', 'Hours/Week', 'Notes']
add_header_row(table12, headers12, '1F4E78')

for i in range(1, 10):
    row = table12.rows[i]
    for j in range(6):
        row.cells[j].text = ''

doc12.save('Templates/12_Assignments_Template.docx')

# ============================================
# 13. Contributions Template
# ============================================
doc13 = Document()
doc13.add_heading('AISCR GLOBAL PMO - VOLUNTEER CONTRIBUTIONS', 0)
doc13.add_paragraph('Volunteer Contributions Template')

table13 = doc13.add_table(rows=10, cols=6)
table13.style = 'Light Grid Accent 1'
headers13 = ['Volunteer', 'Project', 'Task', 'Hours', 'Date', 'Review']
add_header_row(table13, headers13, '1F4E78')

for i in range(1, 10):
    row = table13.rows[i]
    for j in range(6):
        row.cells[j].text = ''

doc13.save('Templates/13_Contributions_Template.docx')

# ============================================
# 14. Portfolio Template
# ============================================
doc14 = Document()
doc14.add_heading('AISCR GLOBAL PMO - PROJECT PORTFOLIO', 0)
doc14.add_paragraph('Project Portfolio Template')

table14 = doc14.add_table(rows=10, cols=8)
table14.style = 'Light Grid Accent 1'
headers14 = ['Project', 'Category', 'Status', 'Start Date', 'End Date', 'PM', 'Analyst', 'Priority']
add_header_row(table14, headers14, '1F4E78')

for i in range(1, 10):
    row = table14.rows[i]
    for j in range(8):
        row.cells[j].text = ''

doc14.add_paragraph('Status Options: Green, Yellow, Red | Priority: 1, 2, 3, 4', style='Intense Quote')

doc14.save('Templates/14_Portfolio_Template.docx')

# ============================================
# 15. Risk Heat Map Template
# ============================================
doc15 = Document()
doc15.add_heading('AISCR GLOBAL PMO - RISK HEAT MAP', 0)
doc15.add_paragraph('Risk Heat Map Template - Score = Likelihood × Impact')

table15 = doc15.add_table(rows=26, cols=3)
table15.style = 'Light Grid Accent 1'
headers15 = ['Likelihood', 'Impact', 'Score']
add_header_row(table15, headers15, '1F4E78')

for i in range(1, 26):
    row = table15.rows[i]
    for j in range(3):
        row.cells[j].text = ''

doc15.save('Templates/15_Risk_Heat_Map_Template.docx')

# ============================================
# 16. Gantt Chart Template
# ============================================
doc16 = Document()
doc16.add_heading('AISCR GLOBAL PMO - GANTT CHART', 0)
doc16.add_paragraph('Gantt Chart Template')

table16 = doc16.add_table(rows=10, cols=5)
table16.style = 'Light Grid Accent 1'
headers16 = ['Task', 'Start Date', 'End Date', 'Duration', 'Status']
add_header_row(table16, headers16, '1F4E78')

for i in range(1, 10):
    row = table16.rows[i]
    for j in range(5):
        row.cells[j].text = ''

doc16.add_paragraph('Status Options: Not Started, In Progress, Complete', style='Intense Quote')

doc16.save('Templates/16_Gantt_Chart_Template.docx')

# ============================================
# 17. Stage-Gate Calculator Template
# ============================================
doc17 = Document()
doc17.add_heading('AISCR GLOBAL PMO - STAGE-GATE CALCULATOR', 0)
doc17.add_paragraph('Stage-Gate Calculator Template - Decision = IF(Score >= 3, "PASS", "FAIL")')

table17 = doc17.add_table(rows=5, cols=3)
table17.style = 'Light Grid Accent 1'
headers17 = ['Gate', 'Score', 'Decision']
add_header_row(table17, headers17, '1F4E78')

data17 = [
    ['Initiation', '', ''],
    ['Planning', '', ''],
    ['Execution', '', ''],
    ['Closure', '', '']
]

for i, row_data in enumerate(data17, 1):
    row = table17.rows[i]
    row.cells[0].text = row_data[0]
    row.cells[1].text = row_data[1]
    row.cells[2].text = row_data[2]

doc17.save('Templates/17_Stage_Gate_Calculator_Template.docx')

# ============================================
# 18. Donor Dashboard Template
# ============================================
doc18 = Document()
doc18.add_heading('AISCR GLOBAL PMO - DONOR DASHBOARD', 0)
doc18.add_paragraph('Donor Dashboard Template')

table18 = doc18.add_table(rows=10, cols=4)
table18.style = 'Light Grid Accent 1'
headers18 = ['Donor', 'Project', 'Funding', 'Status']
add_header_row(table18, headers18, '1F4E78')

for i in range(1, 10):
    row = table18.rows[i]
    for j in range(4):
        row.cells[j].text = ''

doc18.add_paragraph('Status Options: Green, Yellow, Red', style='Intense Quote')

doc18.save('Templates/18_Donor_Dashboard_Template.docx')

# ============================================
# 19. Membership Dashboard Template
# ============================================
doc19 = Document()
doc19.add_heading('AISCR GLOBAL PMO - MEMBERSHIP DASHBOARD', 0)
doc19.add_paragraph('Membership Dashboard Template - Growth % = (Current - Previous) / Previous')

table19 = doc19.add_table(rows=5, cols=4)
table19.style = 'Light Grid Accent 1'
headers19 = ['Member Type', 'Count', 'Growth %', 'Notes']
add_header_row(table19, headers19, '1F4E78')

for i in range(1, 5):
    row = table19.rows[i]
    for j in range(4):
        row.cells[j].text = ''

doc19.save('Templates/19_Membership_Dashboard_Template.docx')

# ============================================
# 20. Volunteer Scorecard Template
# ============================================
doc20 = Document()
doc20.add_heading('AISCR GLOBAL PMO - VOLUNTEER SCORECARD', 0)
doc20.add_paragraph('Volunteer Scorecard Template - Performance Score = (Tasks × 0.4) + (Hours × 0.3) + (Quality × 0.3)')

table20 = doc20.add_table(rows=10, cols=5)
table20.style = 'Light Grid Accent 1'
headers20 = ['Volunteer', 'Tasks Completed', 'Hours', 'Quality Rating', 'Performance Score']
add_header_row(table20, headers20, '1F4E78')

for i in range(1, 10):
    row = table20.rows[i]
    for j in range(5):
        row.cells[j].text = ''

doc20.save('Templates/20_Volunteer_Scorecard_Template.docx')

# ============================================
# 21. PMO KPIs Template
# ============================================
doc21 = Document()
doc21.add_heading('AISCR GLOBAL PMO - PMO KPIs', 0)
doc21.add_paragraph('PMO Key Performance Indicators Template')

table21 = doc21.add_table(rows=8, cols=2)
table21.style = 'Light Grid Accent 1'
headers21 = ['Metric', 'Value']
add_header_row(table21, headers21, '1F4E78')

data21 = [
    ['Total Projects', ''],
    ['Active Projects', ''],
    ['At Risk Projects', ''],
    ['Critical Projects', ''],
    ['Total Funding', ''],
    ['Membership Growth', ''],
    ['Average Volunteer Score', '']
]

for i, row_data in enumerate(data21, 1):
    row = table21.rows[i]
    row.cells[0].text = row_data[0]
    row.cells[1].text = row_data[1]

doc21.save('Templates/21_PMO_KPIs_Template.docx')

# ============================================
# 22. Dashboard Overview Template
# ============================================
doc22 = Document()
doc22.add_heading('AISCR GLOBAL PMO - DASHBOARD OVERVIEW', 0)
doc22.add_paragraph('This template provides an overview structure for the PMO Dashboard.')

p = doc22.add_paragraph()
p.add_run('KPI Cards:').bold = True
p = doc22.add_paragraph('• Total Projects')
p = doc22.add_paragraph('• Active (Green)')
p = doc22.add_paragraph('• At Risk (Yellow)')
p = doc22.add_paragraph('• Critical (Red)')
p = doc22.add_paragraph('• Total Funding')
p = doc22.add_paragraph('• Membership Growth')
p = doc22.add_paragraph('• Average Volunteer Score')

p = doc22.add_paragraph()
p.add_run('Charts:').bold = True
p = doc22.add_paragraph('• Portfolio Status Bar Chart')
p = doc22.add_paragraph('• Donor Funding Bar Chart')
p = doc22.add_paragraph('• Membership Growth Line Chart')
p = doc22.add_paragraph('• Risk Heat Map Snapshot')

doc22.save('Templates/22_Dashboard_Overview_Template.docx')

print("All Word templates created successfully!")

